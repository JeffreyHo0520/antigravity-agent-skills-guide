import os
import sys
import json
import argparse

def extract_pdf(pdf_path, output_json, img_dir):
    import fitz  # PyMuPDF
    
    os.makedirs(img_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    extracted_data = []
    
    print(f"Opening PDF: {pdf_path} (Total pages: {len(doc)})")
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        elements = []
        
        # 1. Extract text blocks
        blocks = page.get_text("blocks")
        for b in blocks:
            x0, y0, x1, y1, text, block_no, block_type = b
            if block_type == 0:  # Text block
                cleaned_text = text.strip()
                if cleaned_text:
                    elements.append({
                        "type": "text",
                        "bbox": [x0, y0, x1, y1],
                        "content": cleaned_text
                    })
                    
        # 2. Extract images
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            xref = img[0]
            rects = page.get_image_rects(xref)
            if not rects:
                bbox = [0, 0, 0, 0]
            else:
                r = rects[0]
                bbox = [r.x0, r.y0, r.x1, r.y1]
                
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                img_filename = f"img_page_{page_num}_{img_idx}.{image_ext}"
                img_filepath = os.path.join(img_dir, img_filename)
                
                with open(img_filepath, "wb") as f:
                    f.write(image_bytes)
                    
                elements.append({
                    "type": "image",
                    "bbox": bbox,
                    "path": os.path.join(img_dir, img_filename),
                    "ext": image_ext
                })
            except Exception as e:
                print(f"Warning: Failed to extract image xref {xref} on page {page_num}: {e}", file=sys.stderr)
                
        # Sort elements by vertical position (y0), then horizontal position (x0)
        elements.sort(key=lambda e: (e["bbox"][1], e["bbox"][0]))
        
        extracted_data.append({
            "page": page_num + 1,
            "elements": elements
        })
        
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(extracted_data, f, ensure_ascii=False, indent=2)
        
    print(f"Success: Extracted text and images to {output_json} and {img_dir}")

def build_docx(json_path, img_dir, output_docx):
    from docx import Document
    from docx.shared import Inches, Pt
    
    if not os.path.exists(json_path):
        print(f"Error: Translated JSON path {json_path} not found.", file=sys.stderr)
        sys.exit(1)
        
    with open(json_path, "r", encoding="utf-8-sig") as f:
        pages = json.load(f)
        
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    print(f"Building Word document: {output_docx} (Total pages: {len(pages)})")
    
    for p_idx, page in enumerate(pages):
        page_num = page["page"]
        elements = page["elements"]
        
        if p_idx > 0:
            doc.add_page_break()
            
        for el in elements:
            if el["type"] == "text":
                text = el["content"]
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                p.add_run(text)
                
            elif el["type"] == "image":
                # Check path; try to resolve relative to working directory or absolute
                img_path = el["path"]
                
                # If image path in JSON is relative to original workspace, check if we need to search it in img_dir
                if not os.path.exists(img_path):
                    img_path_basename = os.path.basename(img_path)
                    img_path = os.path.join(img_dir, img_path_basename)
                    
                if not os.path.exists(img_path):
                    print(f"  Warning: Image {img_path} not found.", file=sys.stderr)
                    continue
                    
                bbox = el["bbox"]
                w_pts = bbox[2] - bbox[0]
                w_in = w_pts / 72.0 if w_pts > 0 else 5.0
                
                if w_in <= 0.1:
                    w_in = 5.0
                elif w_in > 5.5:
                    w_in = 5.5
                    
                try:
                    p = doc.add_paragraph()
                    p.alignment = 1 # Center
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    p.add_run().add_picture(img_path, width=Inches(w_in))
                except Exception as e:
                    print(f"  Warning: Could not insert image {img_path}: {e}", file=sys.stderr)
                    
    doc.save(output_docx)
    print(f"Success: Generated Word document {output_docx}")

def convert_to_pdf(docx_path, output_pdf):
    import win32com.client
    
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(output_pdf)
    
    if not os.path.exists(abs_docx):
        print(f"Error: Word document {docx_path} not found.", file=sys.stderr)
        sys.exit(1)
        
    print(f"Converting {docx_path} to {output_pdf} via MS Word...")
    
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(abs_docx)
        doc.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
        print(f"Success: PDF generated at {output_pdf}")
    except Exception as e:
        print(f"Error: MS Word conversion failed: {e}", file=sys.stderr)
        sys.exit(1)
    finally:
        if doc:
            doc.Close()
        if word:
            word.Quit()

def main():
    parser = argparse.ArgumentParser(description="PDF Translator & Rebuilder CLI")
    subparsers = parser.add_subparsers(dest="command", required=True)
    
    # Extract parser
    parser_extract = subparsers.add_parser("extract", help="Extract text and images from PDF")
    parser_extract.add_argument("--pdf", required=True, help="Path to input PDF file")
    parser_extract.add_argument("--output", required=True, help="Path to output JSON file")
    parser_extract.add_argument("--img-dir", required=True, help="Directory to save extracted images")
    
    # Build parser
    parser_build = subparsers.add_parser("build", help="Build docx from JSON and images")
    parser_build.add_argument("--json", required=True, help="Path to translated JSON file")
    parser_build.add_argument("--img-dir", required=True, help="Directory containing extracted images")
    parser_build.add_argument("--output", required=True, help="Path to output docx file")
    
    # Convert parser
    parser_convert = subparsers.add_parser("convert", help="Convert docx to pdf")
    parser_convert.add_argument("--docx", required=True, help="Path to input docx file")
    parser_convert.add_argument("--output", required=True, help="Path to output pdf file")
    
    args = parser.parse_args()
    
    if args.command == "extract":
        extract_pdf(args.pdf, args.output, args.img-dir if hasattr(args, "img-dir") else args.img_dir)
    elif args.command == "build":
        build_docx(args.json, args.img_dir, args.output)
    elif args.command == "convert":
        convert_to_pdf(args.docx, args.output)

if __name__ == "__main__":
    main()
